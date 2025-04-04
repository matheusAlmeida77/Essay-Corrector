import os
import requests
import openpyxl
import spacy
from openpyxl.chart import AreaChart, Reference, Series
from openpyxl.chart.label import DataLabelList
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt
from tkinter import filedialog
from customtkinter import CTk, CTkButton, CTkEntry, CTkLabel, CTkFrame

def load_spacy_model():
    try:
        nlp = spacy.load('pt_core_news_sm')
    except OSError:
        import subprocess
        subprocess.run(["python", "-m", "spacy", "download", "pt_core_news_sm"])
        nlp = spacy.load('pt_core_news_sm')
    return nlp

# Carregar o modelo de português do SpaCy
nlp = load_spacy_model()

def get_language_tool_corrections(text):
    url = "https://api.languagetool.org/v2/check"
    params = {
        'text': text,
        'language': 'pt-BR'  # Português brasileiro
    }
    response = requests.post(url, data=params)
    if response.status_code == 200:
        return response.json()
    else:
        raise Exception(f"Erro ao acessar a API: {response.status_code}")

def analyze_text(text):
    doc = nlp(text)
    
    adjectives = [token.text for token in doc if token.pos_ == 'ADJ']
    conjunctions = [token.text for token in doc if token.pos_ == 'CCONJ' or token.pos_ == 'SCONJ']
    num_periods = text.count('.')
    num_paragraphs = text.count('\n') + 1

    return {
        'adjectives': adjectives,
        'conjunctions': conjunctions,
        'num_periods': num_periods,
        'num_paragraphs': num_paragraphs,
    }

def save_to_excel(data, filepath, sheet_name="Correções"):
    if os.path.exists(filepath):
        workbook = openpyxl.load_workbook(filepath)
        if sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
        else:
            sheet = workbook.create_sheet(sheet_name)
    else:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = sheet_name

        headers = ["Sala", "Tema da Redação", "Nome do Aluno", "Quantidade de Erros de Gramática", "Quantidade de Erros de Ortografia", "Número de Períodos", "Número de Parágrafos", "Média de Erros"]
        if sheet.max_row == 1:
            sheet.append(headers)

    for entry in data:
        sheet.append([
            entry['class'],
            entry['essay_topic'],
            entry['student_name'],
            len(entry['grammar_errors']),
            len(entry['spelling_errors']),
            entry['num_periods'],
            entry['num_paragraphs'],
            f"=AVERAGE({get_column_letter(sheet.max_column-5)}2:{get_column_letter(sheet.max_column-2)}{sheet.max_row})"
        ])

    generate_charts(workbook, sheet)
    workbook.save(filepath)

def save_text_file(student_name, class_name, essay_topic, original_text, corrections, directory):
    if not os.path.exists(directory):
        os.makedirs(directory)
    filename = os.path.join(directory, f"{student_name}_{essay_topic}.txt")
    with open(filename, "w", encoding="utf-8") as file:
        file.write(f"Texto Original:\n{original_text}\n\n")
        file.write("Correções:\n")
        for correction in corrections:
            file.write(f"{correction['message']} (Sugestão: {', '.join(correction['replacements'])})\n")
            file.write(f"Contexto: {correction['context']['text']}\n\n")

def save_word_file(student_name, class_name, essay_topic, original_text, corrections, directory):
    if not os.path.exists(directory):
        os.makedirs(directory)
    filename = os.path.join(directory, f"{student_name}_{essay_topic}.docx")
    document = Document()
    document.add_heading(f'Redação de {student_name}', 0)

    document.add_heading('Texto Original', level=1)
    document.add_paragraph(original_text)

    document.add_heading('Correções', level=1)
    for correction in corrections:
        p = document.add_paragraph()
        p.add_run(f"{correction['message']} ").bold = True
        p.add_run(f"(Sugestão: {', '.join(correction['replacements'][:3])})")
        p.add_run(f"\nContexto: {correction['context']['text']}\n")

    document.save(filename)

def generate_charts(workbook, sheet):
    # Remover gráficos existentes para evitar duplicação
    sheet._charts = []

    # Adicionar gráfico de área para o número de períodos
    chart1 = AreaChart()
    chart1.title = "Número de Períodos por Aluno"
    chart1.style = 13
    chart1.x_axis.title = 'Aluno'
    chart1.y_axis.title = 'Número de Períodos'
    data = Reference(sheet, min_col=6, min_row=1, max_row=sheet.max_row)
    categories = Reference(sheet, min_col=3, min_row=2, max_row=sheet.max_row)
    series = Series(data, title_from_data=True)
    series.data_labels = DataLabelList()
    series.data_labels.showVal = True
    chart1.series.append(series)
    chart1.set_categories(categories)
    sheet.add_chart(chart1, "I2")

    # Adicionar gráfico de área para o número de parágrafos
    chart2 = AreaChart()
    chart2.title = "Número de Parágrafos por Aluno"
    chart2.style = 13
    chart2.x_axis.title = 'Aluno'
    chart2.y_axis.title = 'Número de Parágrafos'
    data = Reference(sheet, min_col=7, min_row=1, max_row=sheet.max_row)
    categories = Reference(sheet, min_col=3, min_row=2, max_row=sheet.max_row)
    series = Series(data, title_from_data=True)
    series.data_labels = DataLabelList()
    series.data_labels.showVal = True
    chart2.series.append(series)
    chart2.set_categories(categories)
    sheet.add_chart(chart2, "I20")

    # Adicionar gráfico de área para a média de erros
    chart3 = AreaChart()
    chart3.title = "Média de Erros por Aluno"
    chart3.style = 13
    chart3.x_axis.title = 'Aluno'
    chart3.y_axis.title = 'Média de Erros'
    data = Reference(sheet, min_col=8, min_row=1, max_row=sheet.max_row)
    categories = Reference(sheet, min_col=3, min_row=2, max_row=sheet.max_row)
    series = Series(data, title_from_data=True)
    series.data_labels = DataLabelList()
    series.data_labels.showVal = True
    chart3.series.append(series)
    chart3.set_categories(categories)
    sheet.add_chart(chart3, "I38")

def sanitize_directory_name(name):
    # Remover caracteres inválidos e limitar o comprimento do nome do diretório
    valid_chars = "-_.() abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
    sanitized_name = ''.join(c for c in name if c in valid_chars)
    return sanitized_name[:100]  # Limitar a 100 caracteres

def process_file(file_path, class_name, student_name, essay_topic, save_dir):
    # Lê o conteúdo do arquivo selecionado
    _, file_extension = os.path.splitext(file_path)
    if file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            original_text = file.read()
    elif file_extension == '.docx':
        doc = Document(file_path)
        original_text = '\n'.join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Formato de arquivo não suportado.")

    correction_result = get_language_tool_corrections(original_text)
    
    grammar_errors = []
    spelling_errors = []
    corrections = []
    if 'matches' in correction_result:
        for match in correction_result['matches']:
            error_type = "gramática" if match['rule']['issueType'] == 'grammar' else "ortografia"
            if error_type == "gramática":
                grammar_errors.append(f"{match['message']} (Sugestão: {', '.join([r['value'] for r in match['replacements']])})")
            else:
                spelling_errors.append(f"{match['message']} (Sugestão: {', '.join([r['value'] for r in match['replacements']])})")
            corrections.append({
                'message': match['message'],
                'context': match['context'],
                'replacements': [r['value'] for r in match['replacements']] if match['replacements'] else []
            })

    analysis = analyze_text(original_text)
    
    data = [{
        'class': class_name,
        'essay_topic': essay_topic,
        'student_name': student_name,
        'grammar_errors': grammar_errors,
        'spelling_errors': spelling_errors,
        'num_periods': analysis['num_periods'],
        'num_paragraphs': analysis['num_paragraphs'],
    }]

    # Criar a estrutura de pastas
    class_dir = sanitize_directory_name(class_name)
    base_dir = os.path.join(save_dir, class_dir)
    if not os.path.exists(base_dir):
        os.makedirs(base_dir)

    topic_dir = sanitize_directory_name(essay_topic)
    theme_dir = os.path.join(base_dir, topic_dir)
    if not os.path.exists(theme_dir):
        os.makedirs(theme_dir)

    student_dir = sanitize_directory_name(student_name)
    student_dir_path = os.path.join(theme_dir, student_dir)
    if not os.path.exists(student_dir_path):
        os.makedirs(student_dir_path)

    # Salvar as planilhas
    save_to_excel(data, os.path.join(base_dir, "sala.xlsx"), sheet_name="Sala")
    save_to_excel(data, os.path.join(theme_dir, "tema.xlsx"), sheet_name="Tema")
    save_to_excel(data, os.path.join(student_dir_path, "aluno.xlsx"), sheet_name="Aluno")

    # Salvar o arquivo de texto do aluno
    save_text_file(student_name, class_name, essay_topic, original_text, corrections, student_dir_path)

    # Salvar o arquivo Word do aluno
    save_word_file(student_name, class_name, essay_topic, original_text, corrections, student_dir_path)

    print(f"As correções foram salvas nas planilhas 'sala.xlsx', 'tema.xlsx' e 'aluno.xlsx', e nos arquivos '{student_name}_{essay_topic}.txt' e '{student_name}_{essay_topic}.docx'")
    print(f"Os arquivos foram salvos com sucesso no diretório: {save_dir}")

def main():
    def select_file():
        file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt"), ("Word files", "*.docx")])
        if file_path:
            entry_file_path.delete(0, 'end')
            entry_file_path.insert(0, file_path)

    def select_directory():
        directory = filedialog.askdirectory()
        if directory:
            entry_save_dir.delete(0, 'end')
            entry_save_dir.insert(0, directory)

    def process_input():
        file_path = entry_file_path.get()
        class_name = entry_class_name.get()
        student_name = entry_student_name.get()
        essay_topic = entry_essay_topic.get()
        save_dir = entry_save_dir.get()
        if file_path and class_name and student_name and essay_topic and save_dir:
            process_file(file_path, class_name, student_name, essay_topic, save_dir)
        else:
            print("Por favor, preencha todos os campos e selecione um arquivo e um diretório.")

    app = CTk()
    app.title("Analisador de Redações")
    # Definindo o tamanho da janela
    app.geometry("800x600")
    app.minsize(800, 600) # Definindo o tamanho mínimo da janela

    frame = CTkFrame(app)
    frame.pack(pady=20, padx=20, fill='both', expand=True)

    label_file_path = CTkLabel(frame, text="Selecione o arquivo da redação:")
    label_file_path.pack(pady=5)
    entry_file_path = CTkEntry(frame, width=600)
    entry_file_path.pack(pady=5)
    button_file_path = CTkButton(frame, text="Selecionar arquivo", command=select_file)
    button_file_path.pack(pady=5)

    label_save_dir = CTkLabel(frame, text="Selecione o diretório de salvamento:")
    label_save_dir.pack(pady=5)
    entry_save_dir = CTkEntry(frame, width=600)
    entry_save_dir.pack(pady=5)
    button_save_dir = CTkButton(frame, text="Selecionar diretório", command=select_directory)
    button_save_dir.pack(pady=5)

    label_class_name = CTkLabel(frame, text="Sala do aluno:")
    label_class_name.pack(pady=5)
    entry_class_name = CTkEntry(frame, width=600)
    entry_class_name.pack(pady=5)

    label_student_name = CTkLabel(frame, text="Nome do aluno:")
    label_student_name.pack(pady=5)
    entry_student_name = CTkEntry(frame, width=600)
    entry_student_name.pack(pady=5)

    label_essay_topic = CTkLabel(frame, text="Tema da redação:")
    label_essay_topic.pack(pady=5)
    entry_essay_topic = CTkEntry(frame, width=600)
    entry_essay_topic.pack(pady=5)

    button_process = CTkButton(frame, text="Processar Redação", command=process_input)
    button_process.pack(pady=20)

    app.mainloop()

if __name__ == "__main__":
    main()